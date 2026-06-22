# Frappe Assistant Core - AI Assistant integration for Frappe Framework
# Copyright (C) 2025 Paul Clinton
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

"""
File Content Extraction Tool for Data Science Plugin.
Extracts content from various file formats (PDF, images, CSV, Excel, documents) for LLM processing.
"""

import base64
import importlib.util
import io
import json
import mimetypes
import os
import subprocess
import sys
import tempfile
import time
from typing import Any, Dict, Optional

import frappe
from frappe import _

from frappe_assistant_core.core.base_tool import BaseTool

# Cache for cross-interpreter PaddleOCR availability checks, keyed by
# python interpreter path -> (is_available, checked_at_epoch_seconds).
# When PaddleOCR lives in a separate venv (a different Python than the one
# running Frappe), checking availability means spawning that interpreter,
# so we avoid doing it on every single extraction call.
_PADDLE_AVAILABILITY_CACHE: Dict[str, tuple] = {}
_PADDLE_AVAILABILITY_TTL_SECONDS = 300


class ExtractFileContent(BaseTool):
    """
    📄 File Content Extraction Tool for LLM Processing

    Extract content from various file formats and prepare it for LLM analysis
    through MCP tools.

    📁 **SUPPORTED FORMATS**:
    • PDFs - Text extraction, OCR for scanned docs, table extraction
    • Images - OCR text extraction, object detection, content analysis
    • CSV/Excel - Data parsing, validation, transformation, insights
    • Documents - DOCX, TXT content extraction and analysis

    🎯 **KEY CAPABILITIES**:
    • Text extraction from any document type
    • OCR for scanned documents and images
    • Table and structured data extraction
    • Format-aware content parsing
    • Data validation for CSV/Excel
    • Multi-language support for OCR
    • Content preparation for LLM processing

    💡 **USE CASES**:
    • Extract invoice content for LLM analysis
    • Read contracts and legal documents
    • Extract data from scanned forms
    • Parse CSV/Excel data for processing
    • OCR scanned documents
    • Prepare documents for Q&A with LLMs
    • Extract structured data for validation
    """

    def __init__(self):
        super().__init__()
        self.name = "extract_file_content"
        self.description = self._get_description()
        self.requires_permission = "File"  # Requires permission to access File DocType

        self.inputSchema = {
            "type": "object",
            "properties": {
                "file_url": {
                    "type": "string",
                    "description": "File URL from Frappe (e.g., '/files/invoice.pdf' or '/private/files/document.docx'). Provide either file_url OR file_name.",
                },
                "file_name": {
                    "type": "string",
                    "description": "Alternative: File name from File DocType (e.g., 'invoice-2024.pdf'). Provide either file_url OR file_name.",
                },
                "operation": {
                    "type": "string",
                    "enum": ["extract", "ocr", "parse_data", "extract_tables"],
                    "description": "Operation: 'extract' (get text/data), 'ocr' (extract text from images), 'parse_data' (structured data from CSV/Excel), 'extract_tables' (extract tables from PDFs)",
                },
                "language": {
                    "type": "string",
                    "default": "en",
                    "description": "OCR language code (en, fr, german, es, ch, etc.)",
                },
                "output_format": {
                    "type": "string",
                    "enum": ["json", "text", "markdown"],
                    "default": "text",
                    "description": "Output format for extracted content",
                },
                "max_pages": {
                    "type": "integer",
                    "default": 50,
                    "description": "Maximum pages to process for PDFs",
                },
            },
            "required": ["operation"],
        }

    def _get_description(self) -> str:
        """Get tool description"""
        return """Extract text and data from various file formats for analysis and processing. SUPPORTED FORMATS: PDF (text extraction, table extraction), Images JPG/PNG (OCR with PaddleOCR), Spreadsheets CSV/Excel (parse data), Documents DOCX/TXT (text extraction). OPERATIONS: extract (get text content), ocr (optical character recognition on images), parse_data (structured data from CSV/Excel), extract_tables (tables from PDFs). USE CASES: Read invoices, contracts, forms, reports, spreadsheets for analysis and data processing. Requires valid file URL from Frappe file system. Returns extracted content in text or structured format suitable for further processing."""

    def execute(self, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Execute file content extraction"""
        try:
            # Validate dependencies first
            dep_check = self._check_dependencies()
            if not dep_check["success"]:
                return dep_check

            # Get file from Frappe
            file_doc = self._get_file_document(arguments)
            if not file_doc:
                return {"success": False, "error": "File not found or access denied"}

            # Check file size limits
            if not self._check_file_size(file_doc):
                return {"success": False, "error": "File size exceeds limit of 50MB"}

            # Get file content
            file_content = self._get_file_content(file_doc)
            if not file_content:
                return {"success": False, "error": "Failed to read file content"}

            # Detect file type
            file_type = self._detect_file_type(file_doc)

            # Process based on operation
            operation = arguments.get("operation", "extract")

            if operation == "extract":
                result = self._extract_content(file_content, file_type, arguments)
            elif operation == "ocr":
                result = self._perform_ocr(file_content, arguments, file_type=file_type)
            elif operation == "parse_data":
                if file_type in ["csv", "excel"]:
                    result = self._extract_content(file_content, file_type, arguments)
                else:
                    return {
                        "success": False,
                        "error": "parse_data operation only supports CSV and Excel files",
                    }
            elif operation == "extract_tables":
                if file_type == "pdf":
                    result = self._extract_pdf_tables(file_content, arguments)
                else:
                    return {"success": False, "error": "extract_tables operation only supports PDF files"}
            else:
                return {"success": False, "error": f"Unknown operation: {operation}"}

            # Add file metadata to result
            if result.get("success"):
                result["file_info"] = {
                    "name": file_doc.file_name,
                    "type": file_type,
                    "size": file_doc.file_size if hasattr(file_doc, "file_size") else len(file_content),
                    "url": file_doc.file_url,
                }

            return result

        except Exception as e:
            frappe.log_error(title="File Processing Error", message=f"Error processing file: {str(e)}")
            return {"success": False, "error": str(e)}

    def _check_dependencies(self) -> Dict[str, Any]:
        """Check if required dependencies are available"""
        missing_deps = []

        # Determine which backend is configured so we can skip irrelevant checks
        try:
            ocr_settings = self._get_ocr_settings()
            ocr_backend = ocr_settings.get("backend", "paddleocr")
        except Exception:
            ocr_backend = "paddleocr"

        # NOTE: these used to be skipped when ocr_backend == "ollama", on the
        # assumption that Ollama handles everything itself. That's wrong:
        # _perform_ocr() falls back from Ollama -> PaddleOCR -> Tesseract
        # whenever Ollama fails or returns empty text, and BOTH fallbacks
        # need pypdf/Pillow/pytesseract/fitz. Skipping the check meant a
        # missing dependency wasn't caught here — it instead surfaced deep
        # inside the fallback chain, got swallowed by a broad try/except in
        # _try_ollama_ocr(), and the user just saw a generic empty result
        # with no indication anything was actually missing. We now always
        # check every dependency the fallback chain can touch, regardless
        # of which backend is configured.
        try:
            import pypdf
        except ImportError:
            missing_deps.append("pypdf")

        try:
            import PIL
        except ImportError:
            missing_deps.append("Pillow")

        try:
            import pytesseract
        except ImportError:
            missing_deps.append("pytesseract")

        try:
            import fitz  # PyMuPDF — needed to render PDF pages to images for OCR
        except ImportError:
            missing_deps.append("PyMuPDF (pip install pymupdf)")

        try:
            import pandas
        except ImportError:
            missing_deps.append("pandas")

        if missing_deps:
            return {
                "success": False,
                "error": f"Missing dependencies: {', '.join(missing_deps)}. Please install them using pip.",
            }

        return {"success": True}

    def _get_file_document(self, arguments: Dict[str, Any]) -> Optional[Any]:
        """Get file document from Frappe, with parent document authorization."""
        file_url = arguments.get("file_url")
        file_name = arguments.get("file_name")

        try:
            file_doc = None

            if file_url:
                results = frappe.get_all("File", filters={"file_url": file_url}, fields=["*"], limit=1)
                if results:
                    file_doc = frappe.get_doc("File", results[0].name)

            elif file_name:
                results = frappe.get_all("File", filters={"file_name": file_name}, fields=["*"], limit=1)
                if results:
                    file_doc = frappe.get_doc("File", results[0].name)

            if not file_doc:
                return None

            self._check_file_access(file_doc)
            return file_doc

        except frappe.PermissionError:
            raise
        except Exception as e:
            frappe.log_error(f"Error getting file document: {str(e)}")
            return None

    def _check_file_access(self, file_doc) -> None:
        """
        Verify the user has access to this file.

        Access is granted when ANY of these is true:
          1. The file is attached to a document the user can read.
          2. The file is owned by the current user (chat_with_file sets owner=user).
          3. The current user is a System Manager.

        The old code used frappe.only_for("System Manager") for all private
        unattached files. Every file uploaded via chat_with_file is private and
        unattached, so non-admin users always hit a PermissionError, causing
        _get_file_document to return None and extraction to silently return empty.

        Raises:
            frappe.PermissionError: If none of the above conditions is met.
        """
        current_user = frappe.session.user

        # 1. Attached to a document - check parent document permission
        if file_doc.attached_to_doctype and file_doc.attached_to_name:
            if not frappe.has_permission(file_doc.attached_to_doctype, "read", file_doc.attached_to_name):
                frappe.throw(
                    _("You do not have permission to access the document this file is attached to"),
                    frappe.PermissionError,
                )
            return

        # 2 & 3. Private unattached file - owner or System Manager only
        if file_doc.file_url and "/private/" in file_doc.file_url:
            is_owner = (file_doc.owner == current_user)
            is_admin = "System Manager" in frappe.get_roles(current_user)
            if not (is_owner or is_admin):
                frappe.throw(
                    _("You do not have permission to access this private file"),
                    frappe.PermissionError,
                )

    def _check_file_size(self, file_doc) -> bool:
        """Check if file size is within limits"""
        max_size = 50 * 1024 * 1024  # 50MB

        try:
            if hasattr(file_doc, "file_size") and file_doc.file_size:
                return file_doc.file_size <= max_size

            # Try to get file size from file system
            file_path = self._get_file_path(file_doc)
            if file_path and os.path.exists(file_path):
                return os.path.getsize(file_path) <= max_size

            return True  # Allow if we can't determine size

        except Exception:
            return True

    def _get_file_path(self, file_doc) -> Optional[str]:
        """Get absolute file path"""
        if file_doc.file_url:
            if file_doc.file_url.startswith("/private"):
                # Private file
                return frappe.get_site_path(file_doc.file_url.lstrip("/"))
            elif file_doc.file_url.startswith("/files"):
                # Public file
                return frappe.get_site_path("public", file_doc.file_url.lstrip("/"))
        return None

    def _get_file_content(self, file_doc) -> Optional[bytes]:
        """Get file content as bytes — supports local files and S3 (frappe_s3_attachment)."""
        try:
            # 1. Try local filesystem
            file_path = self._get_file_path(file_doc)
            if file_path and os.path.exists(file_path):
                # nosemgrep: frappe-semgrep-rules.rules.security.frappe-security-file-traversal — _get_file_path scopes to /private or /files under the site directory via frappe.get_site_path
                with open(file_path, "rb") as f:
                    return f.read()

            # 2. Try S3 via frappe_s3_attachment
            s3_content = self._get_s3_content(file_doc)
            if s3_content:
                return s3_content

            # 3. Fallback: inline content on the File doc
            if hasattr(file_doc, "content"):
                if isinstance(file_doc.content, str):
                    return base64.b64decode(file_doc.content)
                return file_doc.content

            return None

        except Exception as e:
            frappe.log_error(f"Error reading file content: {str(e)}")
            return None

    def _get_s3_content(self, file_doc) -> Optional[bytes]:
        """Fetch file bytes from S3 if frappe_s3_attachment is installed."""
        try:
            file_url = file_doc.file_url or ""
            if "frappe_s3_attachment" not in file_url:
                return None

            from frappe_s3_attachment.controller import S3Operations

            # S3 key: prefer content_hash, fallback to parsing file_url
            s3_key = (getattr(file_doc, "content_hash", "") or "").strip()
            if not s3_key:
                from urllib.parse import parse_qs, urlparse

                parsed = urlparse(file_url)
                s3_key = parse_qs(parsed.query).get("key", [""])[0]

            if not s3_key:
                return None

            s3_ops = S3Operations()
            response = s3_ops.read_file_from_s3(s3_key)
            return response["Body"].read()

        except ImportError:
            return None
        except Exception as e:
            frappe.log_error(f"S3 file read failed: {str(e)}")
            return None

    def _detect_file_type(self, file_doc) -> str:
        """Detect file type from document"""
        file_name = file_doc.file_name or file_doc.file_url or ""
        file_name_lower = file_name.lower()

        if file_name_lower.endswith(".pdf"):
            return "pdf"
        elif file_name_lower.endswith((".jpg", ".jpeg", ".png", ".bmp", ".tiff")):
            return "image"
        elif file_name_lower.endswith((".csv", ".tsv")):
            return "csv"
        elif file_name_lower.endswith((".xlsx", ".xls")):
            return "excel"
        elif file_name_lower.endswith(".docx"):
            return "docx"
        elif file_name_lower.endswith((".txt", ".text")):
            return "text"
        else:
            # Try to detect from MIME type
            mime_type, _ = mimetypes.guess_type(file_name)
            if mime_type:
                if "pdf" in mime_type:
                    return "pdf"
                elif "image" in mime_type:
                    return "image"
                elif "csv" in mime_type or "tab-separated" in mime_type:
                    return "csv"
                elif "spreadsheet" in mime_type or "excel" in mime_type:
                    return "excel"
                elif "word" in mime_type:
                    return "docx"
                elif "text" in mime_type:
                    return "text"

            return "unknown"

    def _extract_content(
        self, file_content: bytes, file_type: str, arguments: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Extract content based on file type"""
        try:
            if file_type == "pdf":
                return self._extract_pdf_content(file_content, arguments)
            elif file_type == "image":
                return self._extract_image_content(file_content, arguments)
            elif file_type == "csv":
                return self._extract_csv_content(file_content)
            elif file_type == "excel":
                return self._extract_excel_content(file_content)
            elif file_type == "docx":
                return self._extract_docx_content(file_content)
            elif file_type == "text":
                return self._extract_text_content(file_content)
            else:
                return {"success": False, "error": f"Unsupported file type: {file_type}"}

        except Exception as e:
            return {"success": False, "error": f"Content extraction failed: {str(e)}"}

    def _extract_pdf_content(self, file_content: bytes, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Extract content from PDF"""
        try:
            from pypdf import PdfReader

            # Create PDF reader
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)

            max_pages = arguments.get("max_pages", 50)
            num_pages = min(len(reader.pages), max_pages)

            # Extract text from each page
            text_content = []
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")

            combined_text = "\n\n".join(text_content)

            # If no text extracted, this is likely a scanned PDF - auto-fallback to OCR
            if not combined_text.strip():
                return self._perform_ocr(file_content, arguments, file_type="pdf")

            return {
                "success": True,
                "content": combined_text,
                "pages": num_pages,
                "extracted_pages": len(text_content),
            }

        except Exception as e:
            # Don't give up here — a corrupted text layer, an encrypted PDF
            # pypdf can't parse, or even pypdf itself being missing should
            # still get a real shot at OCR before we report failure. The
            # whole point of this tool is "always try every avenue" rather
            # than stopping at the first error.
            ocr_fallback = self._perform_ocr(file_content, arguments, file_type="pdf")
            if ocr_fallback.get("success") and ocr_fallback.get("content", "").strip():
                ocr_fallback["warning"] = f"Text-layer extraction failed ({e}); used OCR instead."
                return ocr_fallback
            return {"success": False, "error": f"PDF extraction error: {str(e)}"}

    def _extract_image_content(self, file_content: bytes, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Extract content from image using OCR"""
        return self._perform_ocr(file_content, arguments, file_type="image")

    def _get_ocr_settings(self) -> Dict[str, Any]:
        """Get OCR backend configuration from Assistant Core Settings."""
        try:
            settings = frappe.get_single("Assistant Core Settings")
            return {
                "backend": getattr(settings, "ocr_backend", "tesseract") or "tesseract",
                "ocr_language": getattr(settings, "ocr_language", "en") or "en",
                "paddleocr_timeout": int(getattr(settings, "paddleocr_timeout", 120) or 120),
                "paddleocr_max_memory_mb": int(getattr(settings, "paddleocr_max_memory_mb", 2048) or 2048),
                # PaddlePaddle doesn't always have wheels for the latest Python
                # (e.g. it currently has no Python 3.14 build). If the bench's
                # own interpreter can't run PaddleOCR, set
                # `paddleocr_python_path` in site_config.json (or via
                # `bench set-config paddleocr_python_path /path/to/venv/bin/python`)
                # to point at a separate venv with an older Python that has
                # paddleocr/paddlepaddle installed. Defaults to the bench's
                # own interpreter when unset.
                "paddleocr_python_path": frappe.conf.get("paddleocr_python_path") or sys.executable,
                "ollama_url": getattr(settings, "ollama_api_url", "http://localhost:11434")
                or "http://localhost:11434",
                "ollama_model": getattr(settings, "ollama_vision_model", "deepseek-ocr:latest")
                or "deepseek-ocr:latest",
                "ollama_timeout": int(getattr(settings, "ollama_request_timeout", 120) or 120),
            }
        except Exception:
            return {"backend": "tesseract", "ocr_language": "en", "paddleocr_python_path": sys.executable}

    def _get_ocr_language(self, arguments: Dict[str, Any], ocr_settings: Dict[str, Any]) -> str:
        """Get OCR language, preferring the per-request argument over settings default."""
        return arguments.get("language") or ocr_settings.get("ocr_language", "en")

    def _is_paddle_ocr_available(self, python_path: Optional[str] = None) -> bool:
        """Check whether PaddleOCR is importable in the interpreter that will
        actually run it.

        If `python_path` is the current interpreter (the common case when
        paddleocr is installed straight into the bench's venv), this is a
        cheap in-process check. If it points at a separate venv (needed when
        the bench's Python is too new for PaddlePaddle's published wheels),
        we spawn that interpreter briefly to ask it — and cache the answer
        for a few minutes since that costs a subprocess start each time.
        """
        python_path = python_path or sys.executable

        if python_path == sys.executable:
            return importlib.util.find_spec("paddleocr") is not None

        cached = _PADDLE_AVAILABILITY_CACHE.get(python_path)
        if cached and (time.time() - cached[1]) < _PADDLE_AVAILABILITY_TTL_SECONDS:
            return cached[0]

        try:
            check = subprocess.run(
                [
                    python_path,
                    "-c",
                    "import importlib.util, sys; sys.exit(0 if importlib.util.find_spec('paddleocr') else 1)",
                ],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=10,
            )
            available = check.returncode == 0
        except Exception:
            # Interpreter path doesn't exist, isn't executable, etc.
            available = False

        _PADDLE_AVAILABILITY_CACHE[python_path] = (available, time.time())
        return available

    def _missing_paddle_ocr_response(self, python_path: Optional[str] = None) -> Dict[str, Any]:
        """Return a clear error when PaddleOCR dependencies are unavailable."""
        python_path = python_path or sys.executable
        return {
            "success": False,
            "error": (
                f"PaddleOCR is not importable using the interpreter at '{python_path}'. "
                "Either install paddleocr and paddlepaddle into that environment, or "
                "set `paddleocr_python_path` in site_config.json to point at a venv "
                "that has them installed (needed if your bench's Python is newer than "
                "PaddlePaddle's published wheels support)."
            ),
            "ocr_backend": "paddleocr",
        }

    def _perform_ocr(
        self, file_content: bytes, arguments: Dict[str, Any], file_type: str = "image"
    ) -> Dict[str, Any]:
        """Perform OCR on image or PDF content.

        Tries the configured backend first (ollama / paddleocr / tesseract).
        Tesseract is then ALWAYS attempted as a guaranteed last-resort fallback
        if the configured backend fails or returns empty — previously this
        safety net only existed on the "ollama" branch; if `ocr_backend` was
        explicitly set to "paddleocr" or "tesseract", a crash/timeout/OOM in
        that single backend meant the whole extraction just failed with
        nothing else attempted. Tesseract is a plain apt package with no
        GPU/model/network dependency, so it's the most reliable thing to
        fall back to no matter what's configured.

        Args:
            file_content: Raw file bytes
            arguments: Tool arguments (language, max_pages, etc.)
            file_type: File type string ("image" or "pdf")
        """
        ocr_settings = self._get_ocr_settings()
        backend = ocr_settings.get("backend")
        paddle_python_path = ocr_settings.get("paddleocr_python_path", sys.executable)

        primary_result = None
        primary_backend_name = backend or "paddleocr"

        if backend == "ollama":
            primary_result = self._try_ollama_ocr(file_content, arguments, file_type, ocr_settings)
        elif backend == "tesseract":
            # Explicit choice — this IS the tesseract attempt, so just return
            # it directly (nothing further to fall back to).
            return self._perform_tesseract_ocr(file_content, arguments, file_type=file_type)
        else:
            # paddleocr (explicit or default)
            if self._is_paddle_ocr_available(paddle_python_path):
                primary_result = self._perform_paddle_ocr(file_content, arguments, file_type, ocr_settings)
            else:
                primary_result = self._missing_paddle_ocr_response(paddle_python_path)

        if (
            primary_result
            and primary_result.get("success")
            and primary_result.get("content", "").strip()
        ):
            return primary_result

        # Primary backend failed or returned empty — guaranteed Tesseract fallback.
        tesseract_result = self._perform_tesseract_ocr(file_content, arguments, file_type=file_type)
        if tesseract_result.get("success") and tesseract_result.get("content", "").strip():
            tesseract_result["warning"] = (
                f"Configured backend '{primary_backend_name}' returned no usable text "
                f"({(primary_result or {}).get('error') or (primary_result or {}).get('message') or 'no detail'}); "
                f"fell back to Tesseract."
            )
            return tesseract_result

        # Nothing worked at all — surface BOTH underlying errors so the real
        # cause is visible instead of a generic "no text found" message.
        primary_detail = (primary_result or {}).get("error") or (primary_result or {}).get("message") or "no result"
        tesseract_detail = tesseract_result.get("error") or tesseract_result.get("message") or "no result"
        return {
            "success": False,
            "error": (
                f"All OCR methods failed to extract text. "
                f"{primary_backend_name}: {primary_detail} | tesseract: {tesseract_detail}"
            ),
            "ocr_backend": primary_backend_name,
        }

    def _perform_paddle_ocr(
        self, file_content: bytes, arguments: Dict[str, Any], file_type: str, ocr_settings: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Perform OCR using PaddleOCR in an isolated subprocess.

        Spawns a child process to run PaddleOCR so that hangs or out-of-memory
        crashes kill only the subprocess, not the Frappe worker. Communicates
        via JSON over stdin/stdout.
        """
        language = self._get_ocr_language(arguments, ocr_settings)
        timeout = ocr_settings.get("paddleocr_timeout", 120)
        max_memory_mb = ocr_settings.get("paddleocr_max_memory_mb", 2048)
        max_pages = arguments.get("max_pages", 50)

        # Advisory memory check — logs a warning but does not block
        self._check_available_memory(max_memory_mb)

        # Write file content to a temp file for the subprocess
        suffix = ".pdf" if file_type == "pdf" else ".png"
        tmp_file = tempfile.NamedTemporaryFile(suffix=suffix, prefix="fac_ocr_", delete=False)
        try:
            if file_type != "pdf":
                # Save image as PNG for consistent handling
                from PIL import Image

                image = Image.open(io.BytesIO(file_content))
                if image.mode in ("RGBA", "P"):
                    image = image.convert("RGB")
                image.save(tmp_file, format="PNG")
            else:
                tmp_file.write(file_content)
            tmp_file.flush()
            tmp_file.close()

            # Build the JSON request for the subprocess
            request_data = json.dumps(
                {
                    "file_path": tmp_file.name,
                    "file_type": file_type,
                    "language": language,
                    "max_pages": max_pages,
                    "max_memory_mb": max_memory_mb,
                }
            )

            # Spawn isolated subprocess
            proc = subprocess.Popen(
                [sys.executable, "-m", "frappe_assistant_core.utils.ocr_subprocess"],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )

            try:
                stdout, stderr = proc.communicate(
                    input=request_data.encode("utf-8"),
                    timeout=timeout,
                )
            except subprocess.TimeoutExpired:
                proc.kill()
                proc.wait()
                frappe.log_error(
                    title="PaddleOCR Timeout",
                    message=f"PaddleOCR subprocess killed after {timeout}s timeout.",
                )
                return {
                    "success": False,
                    "error": (
                        f"PaddleOCR timed out after {timeout} seconds. "
                        "The document may be too large or complex. "
                        "You can increase the timeout in Assistant Core Settings > OCR."
                    ),
                    "ocr_backend": "paddleocr",
                }

            if proc.returncode != 0:
                error_msg = stderr.decode("utf-8", errors="replace").strip()
                frappe.log_error(
                    title="PaddleOCR Subprocess Error",
                    message=f"PaddleOCR subprocess exited with code {proc.returncode}:\n{error_msg[:2000]}",
                )
                # Check for OOM patterns
                if "MemoryError" in error_msg or "Cannot allocate memory" in error_msg:
                    return {
                        "success": False,
                        "error": (
                            f"PaddleOCR ran out of memory (limit: {max_memory_mb}MB). "
                            "The document may be too large. "
                            "You can increase the memory limit in Assistant Core Settings > OCR."
                        ),
                        "ocr_backend": "paddleocr",
                    }
                return {
                    "success": False,
                    "error": f"PaddleOCR failed: {error_msg[:500]}",
                    "ocr_backend": "paddleocr",
                }

            # Parse the JSON result from stdout
            try:
                result = json.loads(stdout.decode("utf-8"))
            except (json.JSONDecodeError, UnicodeDecodeError) as e:
                return {
                    "success": False,
                    "error": f"Failed to parse PaddleOCR output: {str(e)}",
                    "ocr_backend": "paddleocr",
                }

            result["ocr_backend"] = "paddleocr"
            result["ocr_language"] = language
            return result

        finally:
            try:
                os.unlink(tmp_file.name)
            except OSError:
                pass

    def _preprocess_for_ocr(self, image):
        """Clean up a phone/app-scanned page before handing it to Tesseract.

        Phone scans (OKEN Scanner, CamScanner, etc.) are usually low-contrast,
        slightly soft, and full-color when they don't need to be. Converting
        to grayscale + boosting contrast + a light sharpen consistently
        recovers small/blurry text — especially inside dense table cells —
        without needing any extra system dependencies (pure Pillow).
        """
        from PIL import ImageOps, ImageFilter

        if image.mode not in ("L",):
            image = ImageOps.grayscale(image)
        image = ImageOps.autocontrast(image)
        image = image.filter(ImageFilter.SHARPEN)
        return image

    def _tesseract_ocr_image(self, image, language: str) -> str:
        """Run Tesseract on a single preprocessed image, retrying with a
        different page-segmentation mode if the first pass comes back
        suspiciously empty (common on table-heavy or multi-column layouts
        where PSM 3's default layout analysis misreads the structure)."""
        import pytesseract

        processed = self._preprocess_for_ocr(image)

        # PSM 6 ("uniform block of text") tends to do best on invoices/forms
        # with mixed paragraphs + tables, which is the common document type
        # here. PSM 4 ("single column of variable sizes") is tried as a
        # fallback for genuinely columnar/table-dominant layouts.
        text = pytesseract.image_to_string(processed, lang=language, config="--psm 6")
        if len(text.strip()) < 20:
            alt_text = pytesseract.image_to_string(processed, lang=language, config="--psm 4")
            if len(alt_text.strip()) > len(text.strip()):
                text = alt_text
        return text

    def _perform_tesseract_ocr(self, file_content: bytes, arguments: Dict[str, Any], file_type: str = "image") -> Dict[str, Any]:
        """Perform OCR using Tesseract on an image or scanned PDF.

        KEY FIX: Previously used Image.frombytes("RGB", w, h, pix.samples) to
        convert fitz pixmaps to PIL images. This is WRONG when pix.n != 3:
        - pix.n=4 (RGBA, common on PDFs with alpha/transparency in some fitz
          versions): frombytes("RGB") interprets 4-byte RGBA pixels as 3-byte
          RGB, misaligning every pixel after the first. The image looks valid
          (no exception), but is colour-shifted garbage. Tesseract returns
          empty or near-empty output on such a corrupted image.
        - pix.n=1 (grayscale): frombytes("RGB") raises ValueError — but this
          was silently caught by the per-page try/except and logged as a page
          error, never bubbling up as the real root cause.

        The fix: pix.tobytes("png") + Image.open() lets PIL decode the PNG
        with the correct channel count regardless of what fitz produced.
        """
        LANG_MAP = {
            "en": "eng", "fr": "fra", "de": "deu", "es": "spa",
            "it": "ita", "pt": "por", "nl": "nld", "ru": "rus",
            "zh": "chi_sim", "ch": "chi_sim", "ja": "jpn", "ko": "kor",
            "ar": "ara", "hi": "hin", "ta": "tam", "te": "tel",
        }
        ocr_settings = self._get_ocr_settings()
        raw_lang = arguments.get("language") or ocr_settings.get("ocr_language", "en") or "en"
        language = LANG_MAP.get(raw_lang.lower(), raw_lang)

        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            return {
                "success": False,
                "error": "pytesseract or Pillow not installed.",
                "install_command": "pip install pytesseract pillow",
            }

        try:
            pytesseract.get_tesseract_version()
        except Exception:
            return {
                "success": False,
                "error": "Tesseract binary not found. Run: sudo apt-get install tesseract-ocr",
            }

        # PDF: render each page with PyMuPDF then OCR
        if file_type == "pdf":
            try:
                import fitz
            except ImportError:
                return {
                    "success": False,
                    "error": "PyMuPDF not installed. Run: pip install pymupdf",
                }
            try:
                pdf_doc = fitz.open(stream=file_content, filetype="pdf")
            except Exception as e:
                return {"success": False, "error": f"Failed to open PDF: {str(e)}"}

            max_pages = arguments.get("max_pages", 50)
            num_pages = min(len(pdf_doc), max_pages)
            page_texts = []
            page_errors = []

            for page_num in range(num_pages):
                try:
                    page = pdf_doc[page_num]
                    # 400dpi recovers small table text on phone/app-scanned docs
                    # (OKEN Scanner, CamScanner, etc.) that 150-300dpi misses.
                    pix = page.get_pixmap(dpi=400)
                    # FIX: use pix.tobytes("png") + Image.open() instead of
                    # Image.frombytes("RGB", ..., pix.samples).
                    # frombytes("RGB") is WRONG when pix.n != 3:
                    #   n=4 (RGBA): silently produces colour-corrupted image
                    #         → Tesseract returns empty / garbage text
                    #   n=1 (gray): raises ValueError (swallowed by except below)
                    # tobytes("png") always writes a valid PNG with correct
                    # channel count; Image.open() decodes it correctly.
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    if img.mode != "RGB":
                        img = img.convert("RGB")
                    text = self._tesseract_ocr_image(img, language)
                    if text.strip():
                        page_texts.append(f"--- Page {page_num + 1} ---\n{text}")
                    else:
                        # Log per-page empty result so we know Tesseract ran
                        # but found nothing — useful to distinguish from crashes.
                        frappe.log_error(
                            title="Tesseract empty page",
                            message=(
                                f"Page {page_num + 1}: Tesseract ran (lang={language!r}) "
                                f"but returned no text. pix.n={pix.n}, "
                                f"size={pix.width}x{pix.height}."
                            ),
                        )
                except Exception as page_err:
                    page_errors.append(f"Page {page_num + 1}: {str(page_err)}")
                    frappe.log_error(title="Tesseract PDF Page Error", message=f"Page {page_num + 1}: {str(page_err)}")

            pdf_doc.close()
            combined = "\n\n".join(page_texts)

            if not combined.strip():
                if page_errors and len(page_errors) == num_pages:
                    return {
                        "success": False,
                        "error": (
                            f"Tesseract OCR failed on all {num_pages} page(s) — this is a "
                            f"processing error, not an empty document. First error: {page_errors[0]}"
                        ),
                        "ocr_backend": "tesseract",
                    }
                return {"success": True, "content": "", "message": "No text detected in PDF via OCR", "ocr_backend": "tesseract"}

            result = {
                "success": True,
                "content": combined,
                "pages": num_pages,
                "ocr_pages_with_text": len(page_texts),
                "ocr_backend": "tesseract",
                "ocr_language": language,
            }
            if page_errors:
                result["warning"] = (
                    f"{len(page_errors)} of {num_pages} page(s) failed OCR processing "
                    f"and were skipped: {page_errors[0]}"
                )
            return result

        # Image: open directly with PIL
        try:
            image = Image.open(io.BytesIO(file_content))
            if image.mode not in ("RGB",):
                image = image.convert("RGB")
            extracted_text = self._tesseract_ocr_image(image, language)

            if not extracted_text.strip():
                return {"success": True, "content": "", "message": "No text detected in image", "ocr_backend": "tesseract"}

            return {
                "success": True,
                "content": extracted_text,
                "ocr_backend": "tesseract",
                "ocr_language": language,
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Tesseract OCR failed: {str(e)}",
                "ocr_backend": "tesseract",
            }

    def _check_available_memory(self, required_mb: int) -> None:
        """Log a warning if available system memory is below the required threshold.

        Reads /proc/meminfo (Linux only). Advisory only — does not block the OCR call.
        """
        try:
            with open("/proc/meminfo") as f:  # nosemgrep: frappe-security-file-traversal
                for line in f:
                    if line.startswith("MemAvailable:"):
                        available_kb = int(line.split()[1])
                        available_mb = available_kb // 1024
                        if available_mb < required_mb:
                            frappe.log_error(
                                title="PaddleOCR Memory Warning",
                                message=(
                                    f"Low memory before PaddleOCR: {available_mb}MB available, "
                                    f"may need up to {required_mb}MB. "
                                    "OCR will proceed but may fail with out-of-memory error."
                                ),
                            )
                        return
        except Exception:
            pass  # Non-critical; skip on non-Linux or permission errors

    def _try_ollama_ocr(
        self, file_content: bytes, arguments: Dict[str, Any], file_type: str, ocr_settings: Dict[str, Any]
    ) -> Optional[Dict[str, Any]]:
        """Try OCR via Ollama vision model. Returns None on failure to allow PaddleOCR fallback."""
        try:
            if file_type == "pdf":
                return self._perform_ollama_pdf_ocr(file_content, arguments, ocr_settings)
            else:
                # Send image bytes directly as base64 — no PIL required
                return self._ollama_extract_from_bytes(file_content, ocr_settings)
        except Exception as e:
            frappe.log_error(
                title="Ollama OCR Error",
                message=f"Ollama OCR failed: {str(e)}",
            )
            return None

    def _ollama_extract_from_bytes(self, image_bytes: bytes, ocr_settings: Dict[str, Any]) -> Dict[str, Any]:
        """Send raw image bytes to Ollama vision model. No PIL dependency."""
        import requests

        img_b64 = base64.b64encode(image_bytes).decode()

        url = f"{ocr_settings['ollama_url']}/api/generate"
        payload = {
            "model": ocr_settings["ollama_model"],
            "prompt": "Extract all text from this document image exactly as it appears.",
            "images": [img_b64],
            "stream": False,
        }

        response = requests.post(url, json=payload, timeout=ocr_settings["ollama_timeout"])
        response.raise_for_status()
        result = response.json()

        content = result.get("response", "").strip()
        if not content:
            return {"success": True, "content": "", "message": "Ollama returned no text"}
        return {
            "success": True,
            "content": content,
            "ocr_backend": "ollama",
            "ocr_model": ocr_settings["ollama_model"],
        }

    def _ollama_extract_from_image(self, pil_image, ocr_settings: Dict[str, Any]) -> Dict[str, Any]:
        """Send a single PIL image to Ollama vision model for text extraction."""
        import requests

        buf = io.BytesIO()
        if pil_image.mode in ("RGBA", "P"):
            pil_image = pil_image.convert("RGB")
        pil_image.save(buf, format="JPEG", quality=85)
        img_b64 = base64.b64encode(buf.getvalue()).decode()

        url = f"{ocr_settings['ollama_url']}/api/generate"
        payload = {
            "model": ocr_settings["ollama_model"],
            "prompt": "Extract all text from this document image exactly as it appears.",
            "images": [img_b64],
            "stream": False,
        }

        response = requests.post(url, json=payload, timeout=ocr_settings["ollama_timeout"])
        response.raise_for_status()
        result = response.json()

        content = result.get("response", "").strip()
        if not content:
            return {"success": True, "content": "", "message": "Ollama returned no text"}
        return {
            "success": True,
            "content": content,
            "ocr_backend": "ollama",
            "ocr_model": ocr_settings["ollama_model"],
        }

    def _perform_ollama_pdf_ocr(
        self, file_content: bytes, arguments: Dict[str, Any], ocr_settings: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Extract text from PDF via Ollama. Renders pages with PyMuPDF, then sends to Ollama."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            return {
                "success": False,
                "error": "PyMuPDF (fitz) is required for Ollama PDF OCR. Install with: pip install pymupdf",
            }

        try:
            pdf_doc = fitz.open(stream=file_content, filetype="pdf")
        except Exception as e:
            return {"success": False, "error": f"Failed to open PDF for OCR: {str(e)}"}

        max_pages = arguments.get("max_pages", 50)
        num_pages = min(len(pdf_doc), max_pages)

        text_content = []
        for page_num in range(num_pages):
            page = pdf_doc[page_num]
            pix = page.get_pixmap(dpi=150)
            pil_image = pix.pil_image()

            page_result = self._ollama_extract_from_image(pil_image, ocr_settings)
            if page_result.get("success") and page_result.get("content", "").strip():
                text_content.append(f"--- Page {page_num + 1} ---\n{page_result['content']}")

        pdf_doc.close()

        combined_text = "\n\n".join(text_content)

        if not combined_text.strip():
            return {
                "success": True,
                "content": "",
                "message": "Ollama OCR completed but no text was detected.",
                "pages": num_pages,
                "ocr_backend": "ollama",
            }

        return {
            "success": True,
            "content": combined_text,
            "pages": num_pages,
            "ocr_pages_with_text": len(text_content),
            "ocr_backend": "ollama",
            "ocr_model": ocr_settings["ollama_model"],
        }

    def _extract_csv_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from CSV"""
        try:
            import pandas as pd

            # Try different encodings
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    df = pd.read_csv(io.BytesIO(file_content), encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return {"success": False, "error": "Failed to decode CSV file with common encodings"}

            # Convert to dict for serialization
            data_dict = {
                "columns": df.columns.tolist(),
                "row_count": len(df),
                "sample_data": df.head(10).to_dict("records"),
                "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()},
            }

            # Create text representation
            text_content = "CSV Data Summary:\n"
            text_content += f"Columns: {', '.join(data_dict['columns'])}\n"
            text_content += f"Total Rows: {data_dict['row_count']}\n\n"
            text_content += "Sample Data:\n"
            text_content += df.head(10).to_string()

            return {"success": True, "content": text_content, "structured_data": data_dict}

        except Exception as e:
            return {"success": False, "error": f"CSV extraction error: {str(e)}"}

    def _extract_excel_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from Excel"""
        try:
            import pandas as pd

            # Read Excel file
            excel_file = pd.ExcelFile(io.BytesIO(file_content))

            all_sheets_content = []
            structured_data = {}

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)

                # Store structured data
                structured_data[sheet_name] = {
                    "columns": df.columns.tolist(),
                    "row_count": len(df),
                    "sample_data": df.head(10).to_dict("records"),
                }

                # Create text representation
                sheet_content = f"=== Sheet: {sheet_name} ===\n"
                sheet_content += f"Columns: {', '.join(df.columns.tolist())}\n"
                sheet_content += f"Rows: {len(df)}\n\n"
                sheet_content += df.head(10).to_string()

                all_sheets_content.append(sheet_content)

            combined_content = "\n\n".join(all_sheets_content)

            return {
                "success": True,
                "content": combined_content,
                "structured_data": structured_data,
                "sheet_count": len(excel_file.sheet_names),
            }

        except Exception as e:
            return {"success": False, "error": f"Excel extraction error: {str(e)}"}

    def _extract_docx_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from DOCX"""
        try:
            # Check if python-docx is available
            try:
                from docx import Document
            except ImportError:
                return {
                    "success": False,
                    "error": "python-docx not installed. Please install it using: pip install python-docx",
                }

            # Read document
            doc = Document(io.BytesIO(file_content))

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract tables if any
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))

            # Combine content
            content = "\n\n".join(paragraphs)
            if tables_text:
                content += "\n\n=== Tables ===\n\n" + "\n\n".join(tables_text)

            return {
                "success": True,
                "content": content,
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            }

        except Exception as e:
            return {"success": False, "error": f"DOCX extraction error: {str(e)}"}

    def _extract_text_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from text file"""
        try:
            # Try different encodings
            for encoding in ["utf-8", "latin-1", "cp1252", "ascii"]:
                try:
                    text = file_content.decode(encoding)
                    return {"success": True, "content": text, "encoding": encoding}
                except UnicodeDecodeError:
                    continue

            return {"success": False, "error": "Failed to decode text file with common encodings"}

        except Exception as e:
            return {"success": False, "error": f"Text extraction error: {str(e)}"}

    def _extract_pdf_tables(self, file_content: bytes, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Extract tables from PDF"""
        try:
            # Try using pdfplumber for better table extraction
            try:
                import pandas as pd
                import pdfplumber

                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    all_tables = []
                    max_pages = min(arguments.get("max_pages", 50), len(pdf.pages))

                    for page_num in range(max_pages):
                        page = pdf.pages[page_num]
                        tables = page.extract_tables()

                        for table_idx, table in enumerate(tables):
                            if table:
                                # Convert to DataFrame for better structure
                                df = pd.DataFrame(table[1:], columns=table[0] if table else None)
                                all_tables.append(
                                    {
                                        "page": page_num + 1,
                                        "table_index": table_idx + 1,
                                        "data": df.to_dict("records"),
                                        "rows": len(df),
                                        "columns": len(df.columns),
                                    }
                                )

                    if not all_tables:
                        return {"success": True, "message": "No tables found in PDF", "tables": []}

                    return {
                        "success": True,
                        "tables": all_tables,
                        "total_tables": len(all_tables),
                        "pages_processed": max_pages,
                    }

            except ImportError:
                # Fallback to basic extraction if pdfplumber not available
                return {
                    "success": True,
                    "message": "Table extraction requires pdfplumber. Install with: pip install pdfplumber",
                    "fallback": True,
                }

        except Exception as e:
            return {"success": False, "error": f"Table extraction error: {str(e)}"}


# Make sure class is available for discovery
# The plugin manager will find ExtractFileContent automatically